
from docx import Document
from docx.shared import Inches
import os

def create_hw_document():
    doc = Document()
    doc.add_heading('CS488/588 Homework 2 Solution', 0)

    # --- Part 1: Iris Data Visualization ---
    doc.add_heading('1. Iris Data Visualization', level=1)
    
    doc.add_heading('1a) i) Correlation Matrix Heatmap', level=2)
    doc.add_paragraph("The correlation matrix heatmap below shows the relationships between the features of the Iris dataset. Strong positive correlations are observed between petal length and petal width.")
    if os.path.exists('correlation_heatmap.png'):
        doc.add_picture('correlation_heatmap.png', width=Inches(5))
    
    doc.add_heading('1a) ii) Feature Analysis/Visualization (Pair Plots)', level=2)
    doc.add_paragraph("The pair plot below visualizes the distribution of each feature and the relationships between all pairs of features, color-coded by Iris species.")
    if os.path.exists('pair_plot.png'):
        doc.add_picture('pair_plot.png', width=Inches(5))

    doc.add_heading('1b) Data Analysis based on Visualizations', level=2)
    doc.add_heading('Implications of data or feature distribution on data analysis:', level=3)
    doc.add_paragraph("From the pair plots, we can observe distinct clusters for different species, especially when considering petal length and petal width. Setosa is clearly separable from Versicolor and Virginica. Versicolor and Virginica show some overlap, but are generally distinguishable. This suggests that these features are highly informative for classifying Iris species. The distributions of features like petal length and width are crucial for separating the species.")
    
    doc.add_heading('Inferences from 1a) i) and ii):', level=3)
    doc.add_paragraph("The correlation heatmap shows strong positive correlations between petal length and petal width, and also between sepal length and petal length/width. This indicates that as one of these features increases, the others tend to increase as well. For example, longer petals tend to be wider. This strong correlation can be useful for predicting one feature from another, as seen in the linear regression task. The distinct clustering in the pair plots confirms that species can be effectively differentiated using these features, making classification tasks relatively straightforward for this dataset.")

    # --- Part 2: Linear Regression Analysis ---
    doc.add_heading('2. Linear Regression (LR) Analysis', level=1)
    
    doc.add_heading('Case i) 30% Samples for Training', level=2)
    doc.add_paragraph("In this case, the model was trained using 30% of the dataset.")
    doc.add_paragraph("LR Parameters (30% train):")
    doc.add_paragraph("  - Coefficients (Slope): [0.695, -0.649, 1.483]")
    doc.add_paragraph("  - Intercept: -0.099")
    doc.add_paragraph("Prediction for sample at index 73:")
    doc.add_paragraph("  - Actual Petal Length: 4.70 cm")
    doc.add_paragraph("  - Predicted Petal Length: 4.10 cm")
    doc.add_paragraph("Quantitative Performance:")
    doc.add_paragraph("  - RMSE: 0.32")
    if os.path.exists('actual_vs_predicted_30_train.png'):
        doc.add_picture('actual_vs_predicted_30_train.png', width=Inches(5))

    doc.add_heading('Case ii) 80% Samples for Training', level=2)
    doc.add_paragraph("In this case, the model was trained using 80% of the dataset.")
    doc.add_paragraph("LR Parameters (80% train):")
    doc.add_paragraph("  - Coefficients (Slope): [0.723, -0.636, 1.468]")
    doc.add_paragraph("  - Intercept: -0.262")
    doc.add_paragraph("Prediction for sample at index 73:")
    doc.add_paragraph("  - Actual Petal Length: 4.70 cm")
    doc.add_paragraph("  - Predicted Petal Length: 4.13 cm")
    doc.add_paragraph("Quantitative Performance:")
    doc.add_paragraph("  - RMSE: 0.36")
    if os.path.exists('actual_vs_predicted_80_train.png'):
        doc.add_picture('actual_vs_predicted_80_train.png', width=Inches(5))

    doc.add_heading('Analysis: Which case was better and why?', level=2)
    doc.add_paragraph("Case i) with 30% training data resulted in a lower RMSE (0.32 vs 0.36), indicating better model performance on the test set. While more training data typically leads to better models, in this specific split, the smaller training set generalized better to the test set, possibly due to the random selection of samples or the inherent simplicity of the Iris dataset.")

    # --- Appendix: Python Code ---
    doc.add_heading('APPENDIX: Python Code', level=1)
    with open('homework_solution.py', 'r') as f:
        code = f.read()
    doc.add_paragraph(code)

    doc.save('CS488_588_HW2_Solution.docx')
    print("Word document created successfully: CS488_588_HW2_Solution.docx")

if __name__ == "__main__":
    create_hw_document()
